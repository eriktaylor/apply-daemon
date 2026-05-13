"""Tests for the document compiler module."""

import json

import pytest

from src.compile import (
    _apply_executive_summary,
    _extract_clean_bullets,
    _find_resume_baseline,
    _generate_cover_letter,
    generate_assets,
)


@pytest.fixture
def output_dir(tmp_path, monkeypatch):
    monkeypatch.setattr("src.compile.OUTPUT_DIR", tmp_path)
    return tmp_path


@pytest.fixture
def no_resume_baseline(monkeypatch):
    """Ensure no resume baseline is found."""
    monkeypatch.setattr("src.compile._find_resume_baseline", lambda: None)


SAMPLE_CLAUDE_JSON = {
    "match_analysis": "Strong match for this backend role.",
    "custom_cover_letter": "Dear Hiring Manager,\n\nI am excited to apply.\n\nSincerely, Alice",
    "resume_bullet_edits": ["Edit bullet A in Experience section", "Edit bullet B in Skills"],
}

SAMPLE_LISTING = {
    "title": "Senior Backend Engineer",
    "company": "Acme Corp",
    "location": "Remote",
}


class TestGenerateAssets:
    def test_creates_output_directory(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        assert path.exists()
        assert path.is_dir()

    def test_creates_cover_letter_docx(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        cover_letters = list(path.glob("Cover_Letter_*.docx"))
        assert len(cover_letters) == 1
        assert "Acme_Corp" in cover_letters[0].name

    def test_creates_match_analysis_md(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        md = path / "match_analysis.md"
        assert md.exists()
        content = md.read_text()
        assert "Strong match" in content
        assert "Acme Corp" in content

    def test_creates_assets_json(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        assets = json.loads((path / "assets.json").read_text())
        assert assets["match_analysis"] == "Strong match for this backend role."
        assert len(assets["resume_bullet_edits"]) == 2

    def test_directory_name_has_slugs(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        assert "Acme_Corp" in path.name
        assert "Senior_Backend_Engineer" in path.name
        assert "abc12345" in path.name

    def test_creates_edits_only_doc_without_baseline(self, output_dir, no_resume_baseline):
        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        edits_docs = list(path.glob("Resume_Edits_*.docx"))
        assert len(edits_docs) == 1

    def test_creates_targeted_resume_with_baseline(self, output_dir, tmp_path, monkeypatch):
        # Create a fake resume baseline
        from docx import Document
        baseline = tmp_path / "config"
        baseline.mkdir()
        doc = Document()
        doc.add_paragraph("Existing resume content")
        baseline_path = baseline / "My_Resume.docx"
        doc.save(str(baseline_path))
        monkeypatch.setattr("src.compile._find_resume_baseline", lambda: baseline_path)

        path = generate_assets("abc12345-def", SAMPLE_CLAUDE_JSON, SAMPLE_LISTING)
        targeted = list(path.glob("Targeted_Resume_*.docx"))
        assert len(targeted) == 1
        assert "Acme_Corp" in targeted[0].name


class TestGenerateCoverLetter:
    def test_creates_valid_docx(self, tmp_path):
        path = tmp_path / "cover.docx"
        _generate_cover_letter(path, "Dear Hiring Manager,\n\nParagraph 1.\n\nParagraph 2.", "Engineer", "Co")
        assert path.exists()

        from docx import Document
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Dear Hiring Manager" in t for t in texts)

    def test_splits_paragraphs_on_double_newline(self, tmp_path):
        path = tmp_path / "cover.docx"
        _generate_cover_letter(path, "Para one.\n\nPara two.\n\nPara three.", "Engineer", "Co")

        from docx import Document
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) == 3


class TestFindResumeBaseline:
    def test_finds_resume_docx(self, tmp_path, monkeypatch):
        monkeypatch.setattr("src.compile.RESUME_BASELINE_DIR", tmp_path)
        from docx import Document
        doc = Document()
        doc.save(str(tmp_path / "My_Resume_2024.docx"))
        result = _find_resume_baseline()
        assert result is not None
        assert "Resume" in result.name

    def test_returns_none_when_no_resume(self, tmp_path, monkeypatch):
        monkeypatch.setattr("src.compile.RESUME_BASELINE_DIR", tmp_path)
        assert _find_resume_baseline() is None

    def test_prefers_most_recent(self, tmp_path, monkeypatch):
        import time
        monkeypatch.setattr("src.compile.RESUME_BASELINE_DIR", tmp_path)
        from docx import Document
        doc = Document()
        doc.save(str(tmp_path / "old_resume.docx"))
        time.sleep(0.05)
        doc.save(str(tmp_path / "new_resume.docx"))
        result = _find_resume_baseline()
        assert result.name == "new_resume.docx"

    def test_ignores_word_owner_lock_file(self, tmp_path, monkeypatch):
        """Word creates `~$foo.docx` owner-lock files while a doc is open.

        These match the resume glob and have the freshest mtime, so a naive
        `max(..., key=mtime)` picks them and python-docx blows up with
        ``PackageNotFoundError`` trying to parse the lock file as a Zip.
        Regression: !tailor / !regenerate failed with
        "Package not found at 'my_profile/~$se_resume.docx'".
        """
        import time
        monkeypatch.setattr("src.compile.RESUME_BASELINE_DIR", tmp_path)
        from docx import Document
        doc = Document()
        real = tmp_path / "base_resume.docx"
        doc.save(str(real))
        # Simulate Word's lock file appearing later, with a fresher mtime.
        time.sleep(0.05)
        lock = tmp_path / "~$se_resume.docx"
        lock.write_bytes(b"not a real docx")
        result = _find_resume_baseline()
        assert result == real, (
            f"expected real baseline, got {result} (lock file leaked through)"
        )

    def test_ignores_dotfiles(self, tmp_path, monkeypatch):
        """Editor swap/backup dotfiles (e.g. `.#resume.docx`) must not win the glob."""
        monkeypatch.setattr("src.compile.RESUME_BASELINE_DIR", tmp_path)
        from docx import Document
        doc = Document()
        real = tmp_path / "base_resume.docx"
        doc.save(str(real))
        import time
        time.sleep(0.05)
        (tmp_path / ".#resume.docx").write_bytes(b"swap")
        result = _find_resume_baseline()
        assert result == real


class TestExtractCleanBullets:
    """Tests for _extract_clean_bullets — the bifurcation safety layer."""

    def test_structured_edits_return_clean_only(self):
        edits = [
            {
                "original_bullet": "Built APIs",
                "slack_diff": "~~Built APIs~~ **Architected high-throughput REST APIs**",
                "clean_bullet": "Architected high-throughput REST APIs",
            },
        ]
        result = _extract_clean_bullets(edits)
        assert result == ["Architected high-throughput REST APIs"]

    def test_plain_string_fallback(self):
        edits = ["Edit section A: new bullet text", "Edit section B: other text"]
        result = _extract_clean_bullets(edits)
        assert result == edits

    def test_mixed_formats(self):
        edits = [
            {"original_bullet": "x", "slack_diff": "~~x~~ **y**", "clean_bullet": "y"},
            "Legacy plain edit",
        ]
        result = _extract_clean_bullets(edits)
        assert result == ["y", "Legacy plain edit"]

    def test_missing_clean_bullet_skipped(self):
        edits = [{"original_bullet": "x", "slack_diff": "diff only"}]
        result = _extract_clean_bullets(edits)
        assert result == []

    def test_empty_list(self):
        assert _extract_clean_bullets([]) == []


class TestCompilerSafety:
    """Verify that markdown/diff syntax never leaks into .docx output."""

    def test_aggressive_markdown_excluded_from_docx(self, output_dir, no_resume_baseline):
        """Slack diff with bold, strikethrough, and special chars must not
        appear in the .docx — only the clean_bullet text should be present."""
        claude_json = {
            "match_analysis": "Good match.",
            "clean_cover_letter_text": "Dear Hiring Manager, I am writing to apply.",
            "cover_letter_diff_summary": "- ~~Generic opener~~ **Company-specific opener**",
            "resume_bullet_edits": [
                {
                    "original_bullet": "Developed APIs",
                    "slack_diff": "~~Developed APIs~~ **Architected high-throughput ~~REST~~ gRPC APIs** *italic* `code` <html>",
                    "clean_bullet": "Architected high-throughput gRPC APIs for user management",
                },
            ],
        }
        listing = {"title": "Engineer", "company": "TestCo"}
        path = generate_assets("safe12345-xyz", claude_json, listing)

        # Read the resume edits .docx and check content
        from docx import Document
        edits_docs = list(path.glob("Resume_Edits_*.docx"))
        assert len(edits_docs) == 1
        doc = Document(str(edits_docs[0]))
        all_text = " ".join(p.text for p in doc.paragraphs)

        # Clean bullet text SHOULD be present
        assert "Architected high-throughput gRPC APIs" in all_text

        # Markdown/diff syntax MUST NOT be present
        assert "~~" not in all_text
        assert "**" not in all_text
        assert "*italic*" not in all_text
        assert "slack_diff" not in all_text
        assert "<html>" not in all_text

    def test_clean_cover_letter_used_in_docx(self, output_dir, no_resume_baseline):
        """Cover letter .docx must use clean_cover_letter_text, not the diff summary."""
        claude_json = {
            "match_analysis": "Good match.",
            "clean_cover_letter_text": "Dear Hiring Manager, I am writing to apply.\n\nBest regards.",
            "cover_letter_diff_summary": "- ~~Generic~~ **Specific** opener",
        }
        listing = {"title": "Engineer", "company": "TestCo"}
        path = generate_assets("cover12345-xyz", claude_json, listing)

        from docx import Document
        cover_docs = list(path.glob("Cover_Letter_*.docx"))
        assert len(cover_docs) == 1
        doc = Document(str(cover_docs[0]))
        all_text = " ".join(p.text for p in doc.paragraphs)

        assert "Dear Hiring Manager" in all_text
        # Diff summary must NOT be in the cover letter
        assert "~~Generic~~" not in all_text
        assert "**Specific**" not in all_text

    def test_backward_compat_custom_cover_letter(self, output_dir, no_resume_baseline):
        """Legacy custom_cover_letter field should still work as fallback."""
        claude_json = {
            "match_analysis": "OK match.",
            "custom_cover_letter": "Dear team, legacy cover letter text.",
        }
        listing = {"title": "Engineer", "company": "LegacyCo"}
        path = generate_assets("legacy1234-xyz", claude_json, listing)

        from docx import Document
        cover_docs = list(path.glob("Cover_Letter_*.docx"))
        assert len(cover_docs) == 1
        doc = Document(str(cover_docs[0]))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "legacy cover letter text" in all_text

    def test_interview_prep_saved(self, output_dir, no_resume_baseline):
        claude_json = {
            "match_analysis": "Match.",
            "interview_prep_guide": "# Interview Prep\n\nBe ready for system design.",
        }
        listing = {"title": "Engineer", "company": "PrepCo"}
        path = generate_assets("prep12345-xyz", claude_json, listing)
        prep_files = list(path.glob("Interview_Prep_*.md"))
        assert len(prep_files) == 1
        assert "system design" in prep_files[0].read_text()


class TestExecutiveSummaryAndSuggestions:
    """Tests for executive_summary_rewrite and other_suggestions output."""

    def test_executive_summary_saved_as_md(self, output_dir, no_resume_baseline):
        claude_json = {
            "match_analysis": "Strong match.",
            "executive_summary_rewrite": "Experienced AI engineer with a track record of shipping...",
            "resume_bullet_edits": [],
        }
        listing = {"title": "Engineer", "company": "SummaryCo"}
        path = generate_assets("sum12345-xyz", claude_json, listing)
        md = path / "Executive_Summary.md"
        assert md.exists()
        assert "Experienced AI engineer" in md.read_text()

    def test_other_suggestions_saved_as_md(self, output_dir, no_resume_baseline):
        claude_json = {
            "match_analysis": "Good fit.",
            "other_suggestions": "- Add a Skills section\n- Highlight open-source contributions",
        }
        listing = {"title": "Engineer", "company": "SuggCo"}
        path = generate_assets("sug12345-xyz", claude_json, listing)
        md = path / "Resume_Suggestions.md"
        assert md.exists()
        assert "Skills section" in md.read_text()

    def test_missing_fields_do_not_create_files(self, output_dir, no_resume_baseline):
        claude_json = {"match_analysis": "OK."}
        listing = {"title": "Engineer", "company": "NullCo"}
        path = generate_assets("nul12345-xyz", claude_json, listing)
        assert not (path / "Executive_Summary.md").exists()
        assert not (path / "Resume_Suggestions.md").exists()

    def test_apply_executive_summary_heading_match(self):
        """Executive summary is replaced when a summary heading is found."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("Professional Summary")
        doc.add_paragraph("Old generic summary text that spans many many words here.")
        doc.add_paragraph("Experience")

        _apply_executive_summary(doc, "New tailored summary text for target role.")
        texts = [p.text for p in doc.paragraphs]
        assert any("New tailored summary text" in t for t in texts)
        assert not any("Old generic summary" in t for t in texts)

    def test_apply_executive_summary_heuristic_fallback(self):
        """Heuristic replaces the first long paragraph when no heading is found."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Jane Doe · jane@example.com · LinkedIn")
        long_para = " ".join(["word"] * 40)
        doc.add_paragraph(long_para)
        doc.add_paragraph("Experience")

        _apply_executive_summary(doc, "Tailored executive summary.")
        texts = [p.text for p in doc.paragraphs]
        assert "Tailored executive summary." in texts
        assert long_para not in texts  # was replaced

    def test_apply_executive_summary_no_match_leaves_doc_unchanged(self):
        """When no paragraph matches, the doc is not modified."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("Short.")
        original_texts = [p.text for p in doc.paragraphs]

        _apply_executive_summary(doc, "New summary.")
        assert [p.text for p in doc.paragraphs] == original_texts


class TestMissingDiffFailsafe:
    """Verify graceful handling when Claude omits slack_diff keys."""

    def test_format_diff_text_missing_slack_diff(self):
        """When slack_diff is missing, should show 'Diff unavailable' not crash."""
        from src.sweeper import _format_diff_text
        assets_json = {
            "resume_bullet_edits": [
                {"original_bullet": "Built APIs", "clean_bullet": "Architected APIs"},
            ],
        }
        result = _format_diff_text(assets_json)
        assert "Diff unavailable" in result

    def test_format_diff_text_no_edits(self):
        from src.sweeper import _format_diff_text
        assert _format_diff_text({}) == ""

    def test_format_diff_text_with_valid_diffs(self):
        from src.sweeper import _format_diff_text
        assets_json = {
            "resume_bullet_edits": [
                {
                    "original_bullet": "Built APIs",
                    "slack_diff": "~~Built APIs~~ **Architected APIs**",
                    "clean_bullet": "Architected APIs",
                },
            ],
            "cover_letter_diff_summary": "- Added company-specific paragraph",
        }
        result = _format_diff_text(assets_json)
        assert "~~Built APIs~~" in result
        assert "**Architected APIs**" in result
        assert "company-specific paragraph" in result
        assert "Resume Edits" in result
        assert "Cover Letter Updates" in result

    def test_format_diff_text_plain_string_edits(self):
        """Legacy plain-string edits should render without crashing."""
        from src.sweeper import _format_diff_text
        assets_json = {
            "resume_bullet_edits": ["Edit section A: rewrite bullet"],
        }
        result = _format_diff_text(assets_json)
        assert "Edit section A" in result


class TestWritePolishedResumeDocx:
    """_write_polished_resume_docx converts Markdown to styled .docx."""

    def test_inline_bold_produces_bold_run(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, "**Acme Corp** | Senior Engineer | 2020-2023")
        doc = Document(str(dest))
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert any("Acme Corp" in r.text for r in bold_runs)

    def test_inline_italic_produces_italic_run(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, "See *note* below")
        doc = Document(str(dest))
        runs = doc.paragraphs[0].runs
        italic_runs = [r for r in runs if r.italic]
        assert any("note" in r.text for r in italic_runs)

    def test_plain_text_has_no_bold_runs(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, "Plain line with no markup")
        doc = Document(str(dest))
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert not bold_runs

    def test_bullet_with_bold_renders_correctly(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, "- Reduced latency by **40%** via caching")
        doc = Document(str(dest))
        all_text = " ".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "40%" in all_text
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert any("40%" in r.text for r in bold_runs)

    def test_heading_h2_is_bold(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, "## Experience")
        doc = Document(str(dest))
        runs = doc.paragraphs[0].runs
        assert runs[0].bold
        assert "Experience" in runs[0].text

    def test_mixed_document_structure(self, tmp_path):
        from docx import Document

        from src.compile import _write_polished_resume_docx
        md = "\n".join([
            "# Jane Smith",
            "",
            "## Experience",
            "**Acme Corp** | SWE | 2020-2023",
            "- Reduced latency by **40%**",
            "",
            "## Skills",
            "Python, Go",
        ])
        dest = tmp_path / "test.docx"
        _write_polished_resume_docx(dest, md)
        doc = Document(str(dest))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Smith" in all_text
        assert "Experience" in all_text
        assert "Acme Corp" in all_text
        assert "40%" in all_text
