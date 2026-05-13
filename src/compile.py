"""Document compiler — generates .docx application assets from Claude JSON.

Uses python-docx to produce editable Word documents:
  - Cover letter as a new .docx
  - Targeted resume by cloning the baseline and applying bullet edits

Usage:
    from src.compile import generate_assets
    output_dir = generate_assets(job_id, claude_json)
"""

from __future__ import annotations

import json
import logging
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("output")
RESUME_BASELINE_DIR = Path("my_profile")


def _serialize_safe(obj: dict) -> str:
    """JSON-serialize a dict, converting datetime objects to ISO strings."""
    def _default(o: object) -> str:
        if isinstance(o, (datetime, date)):
            return o.isoformat()
        raise TypeError(f"Object of type {type(o).__name__} is not JSON serializable")
    return json.dumps(obj, indent=2, default=_default)


def generate_assets(
    job_id: str,
    claude_json: dict,
    listing: dict,
    *,
    research_context: str = "",
) -> Path:
    """Generate .docx application assets from Claude's JSON response.

    Args:
        job_id: The listing ID.
        claude_json: Parsed JSON with match_analysis, custom_cover_letter, resume_bullet_edits.
        listing: Dict with title, company, etc. for file naming.
        research_context: Raw Deep Research text fed to Claude (for training data).

    Returns:
        Path to the output directory.
    """
    title = listing.get("title", "role")
    company = listing.get("company", "company")
    company_slug = re.sub(r"[^\w]+", "_", company).strip("_")

    title_slug = re.sub(r"[^\w]+", "_", title).strip("_")[:30]
    dir_name = f"{company_slug}_{title_slug}_{job_id[:8]}"
    output_path = OUTPUT_DIR / dir_name
    output_path.mkdir(parents=True, exist_ok=True)

    # Generate cover letter .docx — use clean_cover_letter_text (no markdown),
    # falling back to custom_cover_letter for backward compatibility.
    cover_letter_text = (
        claude_json.get("clean_cover_letter_text")
        or claude_json.get("custom_cover_letter", "")
    )
    if cover_letter_text:
        _generate_cover_letter(
            output_path / f"Cover_Letter_{company_slug}.docx",
            cover_letter_text,
            title,
            company,
        )

    # Generate targeted resume .docx — clone baseline, apply executive summary + bullet edits
    executive_summary = claude_json.get("executive_summary_rewrite", "")
    raw_edits = claude_json.get("resume_bullet_edits", [])
    if raw_edits or executive_summary:
        _generate_targeted_resume(
            output_path,
            company_slug,
            raw_edits,
            executive_summary=executive_summary,
        )

    # Save executive summary rewrite as a standalone Markdown reference
    if executive_summary:
        (output_path / "Executive_Summary.md").write_text(
            f"# Executive Summary Rewrite: {title} at {company}\n\n{executive_summary}\n",
            encoding="utf-8",
        )
        logger.info("Executive summary rewrite saved: Executive_Summary.md")

    # Save free-form resume suggestions
    other_suggestions = claude_json.get("other_suggestions", "")
    if other_suggestions:
        (output_path / "Resume_Suggestions.md").write_text(
            f"# Resume Suggestions: {title} at {company}\n\n{other_suggestions}\n",
            encoding="utf-8",
        )
        logger.info("Resume suggestions saved: Resume_Suggestions.md")

    # Save match analysis as markdown (quick reference)
    (output_path / "match_analysis.md").write_text(
        f"# Match Analysis: {title} at {company}\n\n"
        f"{claude_json.get('match_analysis', '')}\n",
        encoding="utf-8",
    )

    # Save interview prep guide as Markdown if generated
    interview_prep = claude_json.get("interview_prep_guide", "")
    if interview_prep:
        (output_path / f"Interview_Prep_{company_slug}.md").write_text(
            f"# Interview Prep: {title} at {company}\n\n{interview_prep}\n",
            encoding="utf-8",
        )
        logger.info("Interview prep guide saved: Interview_Prep_%s.md", company_slug)

    # Save custom question answers if present (unified intake path)
    custom_answers = claude_json.get("custom_question_answers", [])
    if custom_answers:
        lines = [f"# Custom Application Answers — {company}\n"]
        for i, qa in enumerate(custom_answers, 1):
            q = qa.get("question", "")
            a = qa.get("answer", "")
            lines.append(f"## Q{i}: {q}\n")
            lines.append(f"{a}\n")
        (output_path / "Custom_Answers.md").write_text(
            "\n".join(lines), encoding="utf-8",
        )
        logger.info("Custom answers saved: Custom_Answers.md")

    # Save full JSON for programmatic access
    (output_path / "assets.json").write_text(
        json.dumps(claude_json, indent=2),
        encoding="utf-8",
    )

    # --- Training data dump (DPO flywheel) ---
    # Original triage from local LLM
    (output_path / "original_triage.json").write_text(
        _serialize_safe(listing),
        encoding="utf-8",
    )
    # Full Claude analysis
    (output_path / "tailored_analysis.json").write_text(
        json.dumps(claude_json, indent=2),
        encoding="utf-8",
    )
    # Raw Deep Research source material
    if research_context:
        (output_path / "deep_research_context.txt").write_text(
            research_context,
            encoding="utf-8",
        )

    logger.info("Generated assets in %s", output_path)
    return output_path


def _extract_clean_bullets(raw_edits: list) -> list[str]:
    """Extract clean bullet text from structured or plain-string edits.

    Structured format: {"original_bullet": "...", "slack_diff": "...", "clean_bullet": "..."}
    Plain format: "Edit description text"

    The slack_diff field is intentionally ignored — it contains markdown
    formatting that must never reach the .docx compiler.
    """
    clean: list[str] = []
    for edit in raw_edits:
        if isinstance(edit, dict):
            # Structured diff format — use only clean_bullet
            bullet = edit.get("clean_bullet", "")
            if bullet:
                clean.append(bullet)
        elif isinstance(edit, str):
            # Legacy plain string format
            clean.append(edit)
    return clean


def _generate_cover_letter(path: Path, text: str, title: str, company: str) -> None:
    """Create a cover letter .docx from the provided text."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Split on double newlines for paragraphs, single newlines within a paragraph preserved
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    doc.save(str(path))
    logger.info("Cover letter saved: %s", path.name)


def _generate_targeted_resume(
    output_path: Path,
    company_slug: str,
    raw_edits: list,
    *,
    executive_summary: str = "",
) -> None:
    """Clone the resume baseline and apply executive summary + bullet edits in place.

    Looks for a .docx file in my_profile/ matching *resume* or *Resume*.
    If not found, creates a new doc with the edits as guidance.
    """
    baseline = _find_resume_baseline()

    if baseline:
        doc = Document(str(baseline))
        if executive_summary:
            _apply_executive_summary(doc, executive_summary)
        _apply_bullet_edits(doc, raw_edits)
        out_name = f"Targeted_Resume_{company_slug}.docx"
    else:
        logger.warning(
            "No .docx resume baseline found in %s — creating edits-only doc",
            RESUME_BASELINE_DIR,
        )
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run("Resume Bullet Edits")
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph(
            "No baseline resume .docx was found. Apply these edits manually:"
        )
        clean_edits = _extract_clean_bullets(raw_edits)
        for i, edit in enumerate(clean_edits, 1):
            doc.add_paragraph(f"Edit {i}: {edit}")
        out_name = f"Resume_Edits_{company_slug}.docx"

    save_path = output_path / out_name
    doc.save(str(save_path))
    logger.info("Resume saved: %s", save_path.name)


def _find_resume_baseline() -> Path | None:
    """Find the resume .docx baseline in my_profile/.

    Filters out Microsoft Word owner-lock files (``~$foo.docx``) which Word
    creates next to any open document. Without this filter, when the user
    has their resume open in Word the lock file wins the most-recent-mtime
    tiebreak and python-docx raises ``PackageNotFoundError`` trying to
    parse the lock file as a Zip package.
    """
    candidates = [
        p for p in RESUME_BASELINE_DIR.glob("*[Rr]esume*.docx")
        if not p.name.startswith("~$") and not p.name.startswith(".")
    ]
    if candidates:
        # Prefer the most recently modified
        return max(candidates, key=lambda p: p.stat().st_mtime)
    return None


_SUMMARY_HEADING_KEYWORDS = frozenset({"summary", "profile", "objective", "overview", "about"})


def _apply_executive_summary(doc: Document, new_summary: str) -> None:
    """Replace the executive summary section in the resume .docx.

    Strategy (in order):
    1. Find a heading paragraph that contains a summary keyword, then replace
       the following non-empty paragraph (the summary body text).
    2. Fallback: replace the first paragraph with >30 words in the opening 15
       paragraphs (heuristic catch-all for summary-only resumes).
    If neither heuristic matches, the rewrite is available only via
    Executive_Summary.md and is not applied to the .docx.
    """
    paragraphs = doc.paragraphs

    # Pass 1: heading keyword → replace the next body paragraph
    for i, p in enumerate(paragraphs[:-1]):
        text_lower = p.text.strip().lower()
        if not text_lower:
            continue
        if any(kw in text_lower for kw in _SUMMARY_HEADING_KEYWORDS):
            for j in range(i + 1, min(i + 4, len(paragraphs))):
                next_p = paragraphs[j]
                if next_p.text.strip():
                    next_p.clear()
                    next_p.add_run(new_summary)
                    logger.info(
                        "Executive summary replaced after heading '%s'",
                        p.text.strip()[:40],
                    )
                    return

    # Pass 2: first substantial paragraph (>30 words) in the opening section
    for i, p in enumerate(paragraphs[:15]):
        if len(p.text.strip().split()) > 30:
            p.clear()
            p.add_run(new_summary)
            logger.info("Executive summary replaced at paragraph %d (heuristic)", i)
            return

    logger.info(
        "No executive summary paragraph found in resume — "
        "rewrite saved as Executive_Summary.md only"
    )


def _apply_bullet_edits(doc: Document, raw_edits: list) -> None:
    """Apply bullet edits in place by finding and replacing matching paragraphs.

    Each edit is a dict with ``original_bullet`` and ``clean_bullet`` keys.
    For each edit, scans all paragraphs for a text match against the original
    and replaces the paragraph content while preserving formatting.

    Falls back to plain-string edits (legacy format) by appending them.
    """
    if not raw_edits:
        return

    applied = 0
    unapplied: list[str] = []

    for edit in raw_edits:
        if isinstance(edit, dict):
            original = (edit.get("original_bullet") or "").strip()
            replacement = (edit.get("clean_bullet") or "").strip()
            if not original or not replacement:
                continue

            matched = False
            for p in doc.paragraphs:
                para_text = p.text.strip()
                if not para_text:
                    continue
                # Match if the original is a substring of the paragraph text
                # or if the first 40 characters match (handles minor whitespace diffs)
                if (original in para_text
                        or para_text in original
                        or para_text[:40] == original[:40]):
                    p.clear()
                    p.add_run(replacement)
                    applied += 1
                    matched = True
                    break

            if not matched:
                unapplied.append(replacement)
                logger.debug("No paragraph match for original: %.60s...", original)
        elif isinstance(edit, str):
            # Legacy plain string — can't do in-place replacement
            unapplied.append(edit)

    logger.info("In-place bullet edits: %d applied, %d unmatched", applied, len(unapplied))

    # Append any unmatched edits to the end as a fallback
    if unapplied:
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run("Additional Edits (no matching bullet found)")
        run.bold = True
        run.font.size = Pt(12)
        for i, edit in enumerate(unapplied, 1):
            doc.add_paragraph(f"{i}. {edit}")


def _add_inline_markdown(paragraph, text: str) -> None:
    """Add text to *paragraph* with inline **bold** and *italic* rendered as runs.

    Handles the common resume patterns produced by the polish prompt:
      **bold text**   → bold run
      *italic text*   → italic run
      plain text      → normal run
    """
    # Tokenize on **...** (bold) and *...* (italic), longest match first
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            run = paragraph.add_run(m.group(2))
            run.italic = True
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _write_polished_resume_docx(path: Path, markdown_text: str) -> None:
    """Write a polished resume from Markdown-formatted text into a .docx file.

    Understands Markdown conventions produced by the polish prompt:
      - ``# Name``        → large bold heading
      - ``## Section``    → bold section heading
      - ``- bullet``      → list paragraph
      - plain text        → normal paragraph

    Preserves blank lines as paragraph separators.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in markdown_text.splitlines():
        stripped = line.strip()

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)

        elif stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)

        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_markdown(p, text)

        elif stripped == "":
            # Blank line — spacer only when the last paragraph was non-empty
            paras = doc.paragraphs
            if paras and paras[-1].text.strip():
                doc.add_paragraph("")

        else:
            p = doc.add_paragraph()
            _add_inline_markdown(p, stripped)

    doc.save(str(path))
    logger.info("Polished resume saved: %s", path.name)
